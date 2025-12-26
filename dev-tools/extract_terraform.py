"""
Extrae preguntas del documento Word de Terraform
Maneja texto y preguntas en imágenes
"""
import json
from pathlib import Path
from docx import Document
from docx.oxml.text.paragraph import CT_P
from docx.oxml.table import CT_Tbl
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph

def analyze_docx_structure(docx_path):
    """Analiza la estructura del documento para entender el formato"""
    doc = Document(docx_path)
    
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    print(f"Total tables: {len(doc.tables)}")
    print(f"Total images: {len([r for p in doc.paragraphs for r in p.runs if r._element.xpath('.//pic:pic')])}")
    
    # Mostrar primeros 50 párrafos para entender la estructura
    print("\n" + "="*80)
    print("PRIMEROS PÁRRAFOS:")
    print("="*80)
    for i, para in enumerate(doc.paragraphs[:50]):
        text = para.text.strip()
        if text:
            # Verificar si tiene imágenes
            has_images = bool(para._element.xpath('.//pic:pic'))
            img_marker = " [IMG]" if has_images else ""
            print(f"{i}: {text[:100]}{img_marker}")
    
    # Analizar tablas
    if doc.tables:
        print("\n" + "="*80)
        print("PRIMERA TABLA:")
        print("="*80)
        table = doc.tables[0]
        for i, row in enumerate(table.rows[:10]):
            print(f"Row {i}: {[cell.text[:50] for cell in row.cells]}")

if __name__ == "__main__":
    docx_path = Path(r"c:\Github\OCITest\Descargables\Terraform\HASHICORP TERRAFORM.docx")
    
    if not docx_path.exists():
        print(f"Error: No se encuentra el archivo {docx_path}")
        exit(1)
    
    print(f"Analizando: {docx_path}")
    analyze_docx_structure(docx_path)
